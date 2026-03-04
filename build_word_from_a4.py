from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document


ROOT = Path(__file__).resolve().parent
HTML_PATH = ROOT / "programma-a4.html"
DOCX_PATH = ROOT / "programma-a4.docx"


def clean(text: str) -> str:
    return " ".join(text.replace("\xa0", " ").split())


def main() -> None:
    html = HTML_PATH.read_text(encoding="utf-8")
    soup = BeautifulSoup(html, "html.parser")

    doc = Document()

    h1 = soup.find("h1")
    if h1:
        doc.add_heading(clean(h1.get_text(" ", strip=True)), level=1)

    subtitle = soup.find("p", class_="subtitle")
    if subtitle:
        doc.add_paragraph(clean(subtitle.get_text(" ", strip=True)))

    meta = soup.find("div", class_="meta")
    if meta:
        p = doc.add_paragraph()
        p.add_run(clean(meta.get_text(" ", strip=True)))

    doc.add_paragraph("")

    h2 = soup.find("h2")
    if h2:
        doc.add_heading(clean(h2.get_text(" ", strip=True)), level=2)

    table_html = soup.find("table")
    if table_html:
        rows = table_html.find_all("tr")
        if rows:
            max_cols = max(len(r.find_all(["th", "td"])) for r in rows)
            table = doc.add_table(rows=0, cols=max_cols)
            table.style = "Table Grid"

            for ridx, row in enumerate(rows):
                cells = row.find_all(["th", "td"])
                out_cells = table.add_row().cells
                for cidx, cell in enumerate(cells):
                    txt = clean(cell.get_text(" ", strip=True))
                    out_cells[cidx].text = txt
                    if ridx == 0:
                        for run in out_cells[cidx].paragraphs[0].runs:
                            run.bold = True

    footer_note = soup.find("p", class_="small")
    if footer_note:
        doc.add_paragraph("")
        doc.add_paragraph(clean(footer_note.get_text(" ", strip=True)))

    doc.save(DOCX_PATH)
    print(f"Creato: {DOCX_PATH}")


if __name__ == "__main__":
    main()
